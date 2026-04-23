from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def create_worksheet_docx(title, questions, answers=None, include_answers=False):
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("姓名：__________________  日期：__________________")
    doc.add_paragraph()
    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph(f"{i}. {q}")
        p.paragraph_format.space_after = Pt(12)
    if include_answers and answers:
        doc.add_page_break()
        doc.add_heading("参考答案", level=2)
        for i, ans in enumerate(answers, 1):
            doc.add_paragraph(f"{i}. {ans}")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
